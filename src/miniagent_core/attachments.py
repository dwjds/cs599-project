from __future__ import annotations

import hashlib
import json
import re
from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import unquote
from urllib.request import Request, urlopen
from zipfile import ZipFile

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None  # type: ignore[assignment]

try:
    from openpyxl import Workbook, load_workbook
except ImportError:  # pragma: no cover
    Workbook = None  # type: ignore[assignment]
    load_workbook = None  # type: ignore[assignment]

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None  # type: ignore[assignment]

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas
except ImportError:  # pragma: no cover
    A4 = None  # type: ignore[assignment]
    pdfmetrics = None  # type: ignore[assignment]
    UnicodeCIDFont = None  # type: ignore[assignment]
    canvas = None  # type: ignore[assignment]


def _safe_slug(value: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z_\u4e00-\u9fff.-]+", "_", str(value or "").strip())
    return cleaned.strip("._") or "unknown"


def _safe_filename(filename: str) -> str:
    name = Path(str(filename or "").strip()).name
    name = re.sub(r'[<>:"/\\\\|?*\x00-\x1f]+', "_", name)
    return _safe_slug(name) or "file"


def _safe_stem(filename: str) -> str:
    stem = Path(_safe_filename(filename)).stem
    return stem or "file"


def _sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _guess_kind_from_zip_bytes(data: bytes) -> str | None:
    try:
        with ZipFile(BytesIO(data)) as archive:
            names = set(archive.namelist())
    except Exception:
        return None
    if "word/document.xml" in names:
        return "docx"
    if "xl/workbook.xml" in names:
        return "xlsx"
    return None


def _guess_attachment_kind(
    *,
    filename: str = "",
    content_type: str = "",
    content: bytes | None = None,
    path: Path | None = None,
) -> str | None:
    suffix = Path(filename or (str(path) if path else "")).suffix.lower()
    if suffix in {".pdf", ".docx", ".xlsx", ".txt", ".md", ".py", ".json", ".yaml", ".yml", ".toml", ".csv", ".html", ".css", ".js", ".ts"}:
        return suffix.lstrip(".")

    content_type_lower = (content_type or "").lower()
    if "pdf" in content_type_lower:
        return "pdf"
    if "wordprocessingml.document" in content_type_lower:
        return "docx"
    if "spreadsheetml.sheet" in content_type_lower:
        return "xlsx"
    if content_type_lower.startswith("text/"):
        return "txt"

    raw = content
    if raw is None and path is not None and path.exists():
        raw = path.read_bytes()[:8192]
    if not raw:
        return None
    if raw.startswith(b"%PDF-"):
        return "pdf"
    if raw.startswith(b"PK\x03\x04"):
        return _guess_kind_from_zip_bytes(raw)
    return None


def _ensure_extension(filename: str, kind: str | None) -> str:
    safe_name = _safe_filename(filename)
    if not kind:
        return safe_name
    if Path(safe_name).suffix:
        return safe_name
    return f"{safe_name}.{kind}"


def _attachment_folder_name(filename: str, unique_token: str = "") -> str:
    stem = _safe_stem(filename)
    token = _safe_slug(unique_token)[:8]
    if token:
        return f"{stem}__{token}"
    return stem


def _filename_from_content_disposition(value: str) -> str:
    text = str(value or "").strip()
    if not text:
        return ""
    match = re.search(r"filename\*=([^;]+)", text, flags=re.IGNORECASE)
    if match:
        raw = match.group(1).strip()
        parts = raw.split("''", 1)
        encoded = parts[1] if len(parts) == 2 else raw
        return _safe_filename(unquote(encoded).strip('"'))
    match = re.search(r'filename="?([^";]+)"?', text, flags=re.IGNORECASE)
    if match:
        return _safe_filename(unquote(match.group(1)))
    return ""


def _normalize_text_content(content: str) -> list[str]:
    text = str(content or "").replace("\r\n", "\n")
    return [line.rstrip() for line in text.split("\n")]


def _normalize_table_data(table_data: Any) -> list[list[str]]:
    if table_data is None:
        return []
    rows = table_data
    if isinstance(rows, dict):
        headers = rows.get("headers") or []
        body = rows.get("rows") or []
        normalized: list[list[str]] = []
        if headers:
            normalized.append([str(item) for item in headers])
        for row in body:
            if isinstance(row, dict):
                normalized.append([str(row.get(key, "")) for key in headers])
            else:
                normalized.append([str(item) for item in row])
        return normalized
    normalized_rows: list[list[str]] = []
    for row in rows:
        if isinstance(row, dict):
            if not normalized_rows:
                normalized_rows.append([str(key) for key in row.keys()])
            normalized_rows.append([str(value) for value in row.values()])
        else:
            normalized_rows.append([str(item) for item in row])
    return normalized_rows


@dataclass
class Attachment:
    name: str
    path: str
    source_url: str = ""
    content_type: str = ""
    size: int = 0
    message_id: str = ""
    sha256: str = ""
    origin: str = "inbox"

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


class AttachmentStore:
    """统一管理上传文件(inbox)和结果文件(outbox)。"""

    def __init__(self, workspace: Path):
        self.workspace = workspace
        self.inbox_dir = workspace / "inbox"
        self.outbox_dir = workspace / "outbox"
        self.inbox_dir.mkdir(parents=True, exist_ok=True)
        self.outbox_dir.mkdir(parents=True, exist_ok=True)

    def inbox_sender_dir(self, channel: str, sender_id: str) -> Path:
        return self.inbox_dir / _safe_slug(channel) / _safe_slug(sender_id)

    def outbox_session_dir(self, session_key: str) -> Path:
        return self.outbox_dir / _safe_slug(session_key)

    def _inbox_attachment_dir(
        self,
        channel: str,
        sender_id: str,
        filename: str,
        sha256: str,
    ) -> Path:
        return self.inbox_sender_dir(channel, sender_id) / _attachment_folder_name(
            filename,
            sha256,
        )

    def _outbox_file_dir(self, session_key: str, filename: str) -> Path:
        return self.outbox_session_dir(session_key) / _attachment_folder_name(filename)

    def save_inbound_bytes(
        self,
        *,
        channel: str,
        sender_id: str,
        message_id: str,
        filename: str,
        content: bytes,
        source_url: str = "",
        content_type: str = "",
    ) -> Attachment:
        file_hash = _sha256_bytes(content)
        kind = _guess_attachment_kind(
            filename=filename,
            content_type=content_type,
            content=content[:8192],
        )
        safe_name = _ensure_extension(filename, kind)
        target_dir = self._inbox_attachment_dir(channel, sender_id, safe_name, file_hash)
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / safe_name
        target.write_bytes(content)
        attachment = Attachment(
            name=safe_name,
            path=str(target),
            source_url=source_url,
            content_type=content_type,
            size=len(content),
            message_id=message_id,
            sha256=file_hash,
            origin="inbox",
        )
        existing = self._read_manifest(target_dir)
        existing = [item for item in existing if item.path != attachment.path]
        existing.append(attachment)
        self._write_manifest(target_dir, existing)
        return attachment

    def download_inbound_attachment(
        self,
        *,
        channel: str,
        sender_id: str,
        message_id: str,
        filename: str,
        url: str,
        content_type: str = "",
    ) -> Attachment:
        request = Request(
            url,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
                )
            },
        )
        with urlopen(request, timeout=20) as response:
            content = response.read()
            detected_type = response.headers.get("Content-Type", "") or content_type
            detected_name = _filename_from_content_disposition(
                response.headers.get("Content-Disposition", "")
            )
        resolved_filename = detected_name or filename
        return self.save_inbound_bytes(
            channel=channel,
            sender_id=sender_id,
            message_id=message_id,
            filename=resolved_filename,
            content=content,
            source_url=url,
            content_type=detected_type,
        )

    def save_outbox_text(
        self,
        *,
        session_key: str,
        filename: str,
        content: str,
    ) -> Attachment:
        target_dir = self._outbox_file_dir(session_key, filename)
        target_dir.mkdir(parents=True, exist_ok=True)
        safe_name = _safe_filename(filename)
        target = target_dir / safe_name
        target.write_text(content, encoding="utf-8")
        attachment = Attachment(
            name=safe_name,
            path=str(target),
            content_type="text/plain; charset=utf-8",
            size=len(content.encode("utf-8")),
            sha256=_sha256_bytes(content.encode("utf-8")),
            origin="outbox",
        )
        self._append_manifest(target_dir, attachment)
        return attachment

    def save_outbox_file(
        self,
        *,
        session_key: str,
        filename: str,
        content: str = "",
        title: str = "",
        table_data: Any = None,
        sheet_name: str = "Sheet1",
    ) -> Attachment:
        target_dir = self._outbox_file_dir(session_key, filename)
        target_dir.mkdir(parents=True, exist_ok=True)
        safe_name = _safe_filename(filename)
        target = target_dir / safe_name
        suffix = target.suffix.lower()

        if suffix == ".docx":
            self._write_docx_file(target, title=title, content=content, table_data=table_data)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif suffix == ".pdf":
            self._write_pdf_file(target, title=title, content=content, table_data=table_data)
            content_type = "application/pdf"
        elif suffix == ".xlsx":
            self._write_xlsx_file(
                target,
                title=title,
                content=content,
                table_data=table_data,
                sheet_name=sheet_name,
            )
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        else:
            target.write_text(content, encoding="utf-8")
            content_type = "text/plain; charset=utf-8"

        raw = target.read_bytes()
        attachment = Attachment(
            name=safe_name,
            path=str(target),
            content_type=content_type,
            size=len(raw),
            sha256=_sha256_bytes(raw),
            origin="outbox",
        )
        self._append_manifest(target_dir, attachment)
        return attachment

    def list_session_outbox(self, session_key: str) -> list[Attachment]:
        target_dir = self.outbox_session_dir(session_key)
        if not target_dir.exists():
            return []
        results: list[Attachment] = []
        results.extend(self._read_manifest(target_dir))
        for manifest in target_dir.rglob("manifest.json"):
            if manifest.parent == target_dir:
                continue
            results.extend(self._read_manifest(manifest.parent))
        deduped: list[Attachment] = []
        seen_paths: set[str] = set()
        for item in results:
            normalized = str(Path(item.path).expanduser().resolve())
            if normalized in seen_paths:
                continue
            seen_paths.add(normalized)
            item.path = normalized
            deduped.append(item)
        return deduped

    def collect_session_attachments(
        self,
        messages: list[dict[str, Any]],
        max_items: int = 20,
    ) -> list[Attachment]:
        results: list[Attachment] = []
        seen_paths: set[str] = set()
        for message in reversed(messages):
            for item in reversed(list(message.get("attachments") or [])):
                try:
                    attachment = Attachment(**item)
                except TypeError:
                    continue
                normalized = str(Path(attachment.path).expanduser().resolve())
                if normalized in seen_paths:
                    continue
                if not Path(normalized).exists():
                    continue
                seen_paths.add(normalized)
                attachment.path = normalized
                results.append(attachment)
                if len(results) >= max_items:
                    return results
        return results

    def describe_attachments(self, attachments: list[Attachment]) -> str:
        if not attachments:
            return ""
        lines = [f"用户本轮上传了 {len(attachments)} 个文件："]
        for index, item in enumerate(attachments, start=1):
            details = [f"{index}. {item.name}"]
            if item.content_type:
                details.append(f"type={item.content_type}")
            if item.size:
                details.append(f"size={item.size} bytes")
            details.append(f"path={item.path}")
            lines.append(" | ".join(details))
        return "\n".join(lines)

    def read_text(self, path: str, max_chars: int = 20000) -> str:
        target = Path(path).expanduser().resolve()
        if not target.exists():
            raise FileNotFoundError(f"Attachment not found: {path}")
        kind = _guess_attachment_kind(path=target)

        if kind == "docx":
            return self._read_docx_text(target)[:max_chars]
        if kind == "pdf":
            return self._read_pdf_text(target)[:max_chars]
        if kind == "xlsx":
            return self._read_xlsx_text(target)[:max_chars]
        if kind in {"txt", "md", "py", "json", "yaml", "yml", "toml", "csv", "html", "css", "js", "ts"}:
            return target.read_text(encoding="utf-8")[:max_chars]

        try:
            return target.read_text(encoding="utf-8")[:max_chars]
        except UnicodeDecodeError as exc:
            raise ValueError(
                f"Unsupported non-text attachment format for direct reading: {target.name}"
            ) from exc

    def _read_docx_text(self, path: Path) -> str:
        if DocxDocument is not None:
            document = DocxDocument(path)
            parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts).strip()

        with ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="replace")
        text = re.sub(r"</w:p>", "\n", xml)
        text = re.sub(r"<[^>]+>", "", text)
        compact = re.sub(r"\n{3,}", "\n\n", text)
        return compact.strip()

    def _read_pdf_text(self, path: Path) -> str:
        if PdfReader is None:
            raise ValueError("Reading PDF requires pypdf. Install it with: pip install pypdf")
        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            pages.append((page.extract_text() or "").strip())
        return "\n\n".join(part for part in pages if part).strip()

    def _read_xlsx_text(self, path: Path) -> str:
        if load_workbook is None:
            raise ValueError("Reading XLSX requires openpyxl. Install it with: pip install openpyxl")
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        lines: list[str] = []
        for sheet in workbook.worksheets:
            lines.append(f"[Sheet] {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if item is None else str(item) for item in row]
                if any(values):
                    lines.append(" | ".join(values))
        return "\n".join(lines).strip()

    def _write_docx_file(
        self,
        path: Path,
        *,
        title: str,
        content: str,
        table_data: Any,
    ):
        if DocxDocument is None:
            raise ValueError("Writing DOCX requires python-docx. Install it with: pip install python-docx")
        document = DocxDocument()
        if title:
            document.add_heading(title, level=1)
        for paragraph in _normalize_text_content(content):
            if paragraph.strip():
                document.add_paragraph(paragraph)
        rows = _normalize_table_data(table_data)
        if rows:
            table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    table.cell(row_index, col_index).text = value
        document.save(path)

    def _write_pdf_file(
        self,
        path: Path,
        *,
        title: str,
        content: str,
        table_data: Any,
    ):
        if canvas is None or A4 is None or pdfmetrics is None or UnicodeCIDFont is None:
            raise ValueError("Writing PDF requires reportlab. Install it with: pip install reportlab")
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        page = canvas.Canvas(str(path), pagesize=A4)
        width, height = A4
        y = height - 48
        page.setFont("STSong-Light", 12)

        def write_line(line: str):
            nonlocal y
            if y < 48:
                page.showPage()
                page.setFont("STSong-Light", 12)
                y = height - 48
            page.drawString(48, y, line[:90])
            y -= 18

        if title:
            page.setFont("STSong-Light", 16)
            write_line(title)
            page.setFont("STSong-Light", 12)
            y -= 4

        for paragraph in _normalize_text_content(content):
            write_line(paragraph or " ")

        rows = _normalize_table_data(table_data)
        if rows:
            write_line(" ")
            for row in rows:
                write_line(" | ".join(row))

        page.save()

    def _write_xlsx_file(
        self,
        path: Path,
        *,
        title: str,
        content: str,
        table_data: Any,
        sheet_name: str,
    ):
        if Workbook is None:
            raise ValueError("Writing XLSX requires openpyxl. Install it with: pip install openpyxl")
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = (sheet_name or "Sheet1")[:31]

        row_index = 1
        if title:
            sheet.cell(row=row_index, column=1, value=title)
            row_index += 2

        rows = _normalize_table_data(table_data)
        if rows:
            for row in rows:
                for col_index, value in enumerate(row, start=1):
                    sheet.cell(row=row_index, column=col_index, value=value)
                row_index += 1
        else:
            for line in _normalize_text_content(content):
                sheet.cell(row=row_index, column=1, value=line)
                row_index += 1

        workbook.save(path)

    def _manifest_path(self, directory: Path) -> Path:
        return directory / "manifest.json"

    def _write_manifest(self, directory: Path, attachments: list[Attachment]):
        manifest_path = self._manifest_path(directory)
        manifest_path.write_text(
            json.dumps([item.to_dict() for item in attachments], ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    def _append_manifest(self, directory: Path, attachment: Attachment):
        existing = self._read_manifest(directory)
        existing = [item for item in existing if item.path != attachment.path]
        existing.append(attachment)
        self._write_manifest(directory, existing)

    def _read_manifest(self, directory: Path) -> list[Attachment]:
        manifest_path = self._manifest_path(directory)
        if not manifest_path.exists():
            return []
        payload = json.loads(manifest_path.read_text(encoding="utf-8"))
        results: list[Attachment] = []
        for item in payload:
            results.append(Attachment(**item))
        return results
