from __future__ import annotations

import argparse
import asyncio
import json
import time
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from .attachments import Attachment, AttachmentStore
from .config import EMBEDDING_MODEL, MODEL, WORKSPACE, client
from .harness.trace import ToolEvent, classify_tool_failure
from .memory import MemoryItem, MemoryStore, _memory_item_id


DEFAULT_TASKS_FILE = WORKSPACE / "benchmarks" / "tasks.json"
DEFAULT_MEMORY_TASKS_FILE = WORKSPACE / "benchmarks" / "memory_retrieval_tasks.json"
DEFAULT_RESULTS_DIR = WORKSPACE / "benchmarks" / "results"
FIXTURES_DIR = WORKSPACE / "benchmarks" / "fixtures"


@dataclass
class BenchmarkTask:
    id: str
    category: str
    prompt: str
    attachments: list[dict[str, str]] = field(default_factory=list)
    expected_reply_contains: list[str] = field(default_factory=list)
    expected_reply_contains_any: list[str] = field(default_factory=list)
    expected_tools_all: list[str] = field(default_factory=list)
    expected_tools_any: list[str] = field(default_factory=list)
    expected_outbox_suffixes: list[str] = field(default_factory=list)
    expected_outbox_files: list[dict[str, Any]] = field(default_factory=list)
    expected_max_tool_calls: int | None = None
    max_iterations: int = 10


@dataclass
class BenchmarkResult:
    task_id: str
    category: str
    success: bool
    failure_types: list[str]
    reply_preview: str
    tool_call_count: int
    tool_call_batches: int
    steps: int
    duration_seconds: float
    tool_names: list[str]
    outbox_files: list[str]
    finish_reason: str
    tool_events: list[ToolEvent]


@dataclass
class MemoryRetrievalCase:
    id: str
    query: str
    relevant_ids: list[str]
    top_k: int = 4
    candidate_pool: int = 8


@dataclass
class MemoryRetrievalResult:
    case_id: str
    query: str
    relevant_ids: list[str]
    retrieved_ids: list[str]
    hit: bool
    rank: int | None
    reciprocal_rank: float
    top_k: int


def classify_loop_error(error: str) -> str:
    lowered = str(error or "").lower()
    if "permissiondenied" in lowered or "403" in lowered or "quota" in lowered:
        return "llm_permission_or_quota_error"
    if "timeout" in lowered:
        return "llm_timeout"
    if "connection" in lowered or "network" in lowered:
        return "llm_network_error"
    return "agent_loop_exception"


async def run_benchmark(
    *,
    tasks_file: Path = DEFAULT_TASKS_FILE,
    results_dir: Path = DEFAULT_RESULTS_DIR,
    limit: int | None = None,
    delay_seconds: float = 3.0,
    model: str = MODEL,
    harness: Any | None = None,
) -> dict[str, Any]:
    ensure_default_benchmark_files(tasks_file)
    tasks = load_tasks(tasks_file)
    if limit is not None:
        tasks = tasks[: max(0, limit)]

    if harness is None:
        from miniagent_core.harness import HarnessConfig, MiniAgentHarness

        harness = MiniAgentHarness(HarnessConfig(model=model))

    run_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    results_dir.mkdir(parents=True, exist_ok=True)
    task_results: list[BenchmarkResult] = []

    for index, task in enumerate(tasks, start=1):
        if index > 1 and delay_seconds > 0:
            print(f"[Benchmark] Waiting {delay_seconds:g}s before next task...")
            await asyncio.sleep(delay_seconds)
        print(f"[Benchmark] ({index}/{len(tasks)}) {task.id}: {task.prompt}")
        result = await run_task(task, run_id=run_id, model=model, harness=harness)
        task_results.append(result)
        status = "PASS" if result.success else "FAIL"
        print(
            f"[Benchmark] {status} {task.id} "
            f"steps={result.steps} tools={result.tool_call_count} failures={result.failure_types}"
        )

    summary = build_summary(task_results, run_id=run_id, tasks_file=tasks_file)
    report = {
        "summary": summary,
        "results": [
            {
                **asdict(result),
                "tool_events": [asdict(event) for event in result.tool_events],
            }
            for result in task_results
        ],
    }
    json_path = results_dir / f"{run_id}.json"
    md_path = results_dir / "latest.md"
    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path.write_text(render_markdown_report(summary, task_results), encoding="utf-8")
    report["summary"]["json_report"] = str(json_path)
    report["summary"]["markdown_report"] = str(md_path)
    return report


async def run_task(
    task: BenchmarkTask,
    *,
    run_id: str,
    model: str,
    harness: Any | None = None,
) -> BenchmarkResult:
    start = time.perf_counter()
    if harness is None:
        from miniagent_core.harness import HarnessConfig, MiniAgentHarness

        harness = MiniAgentHarness(HarnessConfig(model=model))

    session = harness.build_eval_session(
        run_id=run_id,
        task_id=task.id,
    )
    store = session.components.attachment_store
    attachments = prepare_task_attachments(task, store, run_id=run_id)
    prompt_content = build_task_prompt(task.prompt, attachments, store)

    turn = await session.run_turn(
        prompt_content,
        attachments=attachments,
        max_iterations=task.max_iterations,
    )
    failure_types = evaluate_task(
        task,
        turn.reply,
        turn.tool_events,
        turn.new_outbox_files,
        turn.metrics,
    )
    if turn.loop_error:
        failure_types = dedupe([*failure_types, classify_loop_error(turn.loop_error)])
    duration = time.perf_counter() - start
    session.components.trace_sink.write(
        "judge_result",
        task_id=task.id,
        category=task.category,
        success=not failure_types,
        failure_types=failure_types,
        duration_seconds=round(duration, 4),
    )
    return BenchmarkResult(
        task_id=task.id,
        category=task.category,
        success=not failure_types,
        failure_types=failure_types,
        reply_preview=str(turn.reply or "")[:1000],
        tool_call_count=int(turn.metrics.get("tool_calls", len(turn.tool_events))),
        tool_call_batches=int(turn.metrics.get("tool_call_batches", 0)),
        steps=int(turn.metrics.get("iterations", 0)),
        duration_seconds=round(duration, 4),
        tool_names=[event.tool for event in turn.tool_events],
        outbox_files=turn.new_outbox_files,
        finish_reason=str(turn.metrics.get("finish_reason", "")),
        tool_events=turn.tool_events,
    )


def build_task_prompt(prompt: str, attachments: list[Attachment], store: AttachmentStore) -> str:
    if not attachments:
        return prompt
    attachment_note = store.describe_attachments(attachments)
    return (
        f"{attachment_note}\n\n"
        "如需查看上传文件内容，请使用 `list_uploaded_files` 或 `read_uploaded_file`。"
        "如果需要输出结果文件，请使用 `save_outbox_file` 保存到 workspace/outbox。\n\n"
        f"用户要求：{prompt}\n\n"
        "回复要求：除非用户明确要求全文或完整展开，否则仅提供简洁摘要，"
        "优先返回关键结果、产物路径和必要验证信息。"
    )


def prepare_task_attachments(task: BenchmarkTask, store: AttachmentStore, *, run_id: str) -> list[Attachment]:
    attachments: list[Attachment] = []
    for index, spec in enumerate(task.attachments, start=1):
        fixture_name = spec.get("fixture", "")
        if not fixture_name:
            continue
        fixture_path = FIXTURES_DIR / fixture_name
        if not fixture_path.exists():
            raise FileNotFoundError(f"Benchmark fixture not found: {fixture_path}")
        filename = spec.get("filename") or fixture_path.name
        attachments.append(
            store.save_inbound_bytes(
                channel="benchmark",
                sender_id=task.id,
                message_id=f"{run_id}-{task.id}-{index}",
                filename=filename,
                content=fixture_path.read_bytes(),
                content_type=spec.get("content_type", ""),
            )
        )
    return attachments


def evaluate_task(
    task: BenchmarkTask,
    reply: str,
    events: list[ToolEvent],
    outbox_files: list[str],
    metrics: dict[str, Any],
) -> list[str]:
    failure_types: list[str] = []
    reply_text = str(reply or "")
    reply_lower = reply_text.lower()
    tool_names = [event.tool for event in events]

    if client is None:
        failure_types.append("llm_unavailable")
    if reply_text.startswith("Error:"):
        failure_types.append("final_error")
    if "Max iterations reached" in reply_text:
        failure_types.append("max_iterations")
    if str(metrics.get("finish_reason", "")) in {"forced_tool_use_failed", "incomplete_progress", "max_iterations"}:
        failure_types.append(str(metrics["finish_reason"]))

    for expected in task.expected_reply_contains:
        if str(expected).lower() not in reply_lower:
            failure_types.append("missing_expected_reply")
            break

    if task.expected_reply_contains_any and not any(
        str(expected).lower() in reply_lower for expected in task.expected_reply_contains_any
    ):
        failure_types.append("missing_expected_reply")

    for expected_tool in task.expected_tools_all:
        if expected_tool not in tool_names:
            failure_types.append("missing_expected_tool")
            break

    if task.expected_tools_any and not any(tool in tool_names for tool in task.expected_tools_any):
        failure_types.append("missing_expected_tool")

    if task.expected_max_tool_calls is not None and len(events) > task.expected_max_tool_calls:
        failure_types.append("too_many_tool_calls")

    for suffix in task.expected_outbox_suffixes:
        normalized_suffix = str(suffix).lower()
        if not any(Path(path).suffix.lower() == normalized_suffix for path in outbox_files):
            failure_types.append("missing_expected_outbox")
            break

    for expectation in task.expected_outbox_files:
        matched, file_failures = evaluate_outbox_file_expectation(expectation, outbox_files)
        if not matched:
            failure_types.extend(file_failures or ["missing_expected_outbox_file"])

    blocking_tool_failures = [
        event.failure_type
        for event in events
        if event.failed and event.failure_type in {"tool_not_found", "invalid_tool_arguments", "skill_script_not_found"}
    ]
    if blocking_tool_failures:
        failure_types.extend(blocking_tool_failures)

    metric_tool_failures = []
    for item in metrics.get("tool_errors", []) or []:
        if not isinstance(item, dict):
            continue
        _, failure_type = classify_tool_failure(str(item.get("tool") or ""), str(item.get("preview") or ""))
        if failure_type in {"tool_not_found", "invalid_tool_arguments", "skill_script_not_found"}:
            metric_tool_failures.append(failure_type)
    if metric_tool_failures:
        failure_types.extend(metric_tool_failures)

    return dedupe(failure_types)


def evaluate_outbox_file_expectation(
    expectation: dict[str, Any],
    outbox_files: list[str],
) -> tuple[bool, list[str]]:
    suffix = str(expectation.get("suffix") or "").lower()
    candidates = [
        Path(path)
        for path in outbox_files
        if not suffix or Path(path).suffix.lower() == suffix
    ]
    if not candidates:
        return False, ["missing_expected_outbox_file"]

    failures: list[str] = []
    for path in candidates:
        try:
            content = read_generated_file_for_assertion(path, expectation)
        except Exception as exc:
            failures.append(f"outbox_content_read_error:{type(exc).__name__}")
            continue
        file_failures = assert_generated_content(content, expectation)
        if not file_failures:
            return True, []
        failures.extend(file_failures)
    return False, dedupe(failures or ["outbox_content_mismatch"])


def read_generated_file_for_assertion(path: Path, expectation: dict[str, Any]) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json", ".csv"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        return {"text": text, "rows": [], "sheet_names": []}
    if suffix == ".docx":
        return {"text": read_docx_text(path), "rows": [], "sheet_names": []}
    if suffix == ".pdf":
        return {"text": read_pdf_text(path), "rows": [], "sheet_names": []}
    if suffix == ".xlsx":
        return read_xlsx_content(path, str(expectation.get("sheet") or ""))
    return {"text": path.read_text(encoding="utf-8", errors="replace"), "rows": [], "sheet_names": []}


def assert_generated_content(content: dict[str, Any], expectation: dict[str, Any]) -> list[str]:
    failures: list[str] = []
    text = str(content.get("text") or "").lower()
    rows = list(content.get("rows") or [])
    sheet_names = [str(name) for name in content.get("sheet_names") or []]

    contains_all = [str(item).lower() for item in expectation.get("contains_all", [])]
    for expected in contains_all:
        if expected not in text:
            failures.append("outbox_missing_text")
            break

    contains_any = [str(item).lower() for item in expectation.get("contains_any", [])]
    if contains_any and not any(expected in text for expected in contains_any):
        failures.append("outbox_missing_any_text")

    not_contains = [str(item).lower() for item in expectation.get("not_contains", [])]
    if any(unexpected in text for unexpected in not_contains):
        failures.append("outbox_unexpected_text")

    expected_sheet = str(expectation.get("sheet") or "")
    if expected_sheet and expected_sheet not in sheet_names:
        failures.append("outbox_missing_sheet")

    min_rows = expectation.get("min_rows")
    if min_rows is not None and len(rows) < int(min_rows):
        failures.append("outbox_too_few_rows")

    max_rows = expectation.get("max_rows")
    if max_rows is not None and len(rows) > int(max_rows):
        failures.append("outbox_too_many_rows")

    return dedupe(failures)


def read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX benchmark assertions") from exc
    document = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            values = [cell.text for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required for PDF benchmark assertions") from exc
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def read_xlsx_content(path: Path, sheet_name: str = "") -> dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("openpyxl is required for XLSX benchmark assertions") from exc
    workbook = load_workbook(str(path), read_only=True, data_only=False)
    sheet = workbook[sheet_name] if sheet_name and sheet_name in workbook.sheetnames else workbook.active
    rows: list[list[Any]] = []
    text_parts: list[str] = []
    for row in sheet.iter_rows(values_only=True):
        values = ["" if value is None else value for value in row]
        if not any(str(value).strip() for value in values):
            continue
        rows.append(values)
        text_parts.append(" | ".join(str(value) for value in values))
    return {
        "text": "\n".join(text_parts),
        "rows": rows,
        "sheet_names": list(workbook.sheetnames),
    }


def build_summary(
    results: list[BenchmarkResult],
    *,
    run_id: str,
    tasks_file: Path,
) -> dict[str, Any]:
    total = len(results)
    passed = sum(1 for result in results if result.success)
    failed = total - passed
    tool_calls = [result.tool_call_count for result in results]
    steps = [result.steps for result in results]
    failure_counts: dict[str, int] = {}
    for result in results:
        for failure_type in result.failure_types:
            failure_counts[failure_type] = failure_counts.get(failure_type, 0) + 1
    return {
        "run_id": run_id,
        "tasks_file": str(tasks_file),
        "total": total,
        "passed": passed,
        "failed": failed,
        "success_rate": round(passed / total, 4) if total else 0.0,
        "total_tool_calls": sum(tool_calls),
        "avg_tool_calls": round(sum(tool_calls) / total, 3) if total else 0.0,
        "avg_steps": round(sum(steps) / total, 3) if total else 0.0,
        "failure_types": failure_counts,
    }


async def run_memory_retrieval_benchmark(
    *,
    tasks_file: Path = DEFAULT_MEMORY_TASKS_FILE,
    results_dir: Path = DEFAULT_RESULTS_DIR,
) -> dict[str, Any]:
    ensure_default_memory_retrieval_tasks(tasks_file)
    payload = json.loads(tasks_file.read_text(encoding="utf-8"))
    memory_items = build_memory_items_from_payload(payload)
    cases = [
        MemoryRetrievalCase(**item)
        for item in payload.get("queries", [])
        if isinstance(item, dict)
    ]

    run_id = datetime.now().strftime("%Y%m%d_%H%M%S")
    workspace = WORKSPACE / "benchmarks" / "tmp" / f"memory_retrieval_{run_id}"
    store = MemoryStore(workspace)
    store.upsert_memory_items(memory_items)
    results: list[MemoryRetrievalResult] = []

    for case in cases:
        retrieved = await store.retrieve_relevant_memory(
            case.query,
            client=client,
            embedding_model=EMBEDDING_MODEL,
            top_k=case.top_k,
            candidate_pool=case.candidate_pool,
        )
        retrieved_ids = [item.id for item in retrieved]
        rank = first_relevant_rank(retrieved_ids, case.relevant_ids)
        results.append(
            MemoryRetrievalResult(
                case_id=case.id,
                query=case.query,
                relevant_ids=list(case.relevant_ids),
                retrieved_ids=retrieved_ids,
                hit=rank is not None,
                rank=rank,
                reciprocal_rank=(1 / rank) if rank else 0.0,
                top_k=case.top_k,
            )
        )

    summary = build_memory_retrieval_summary(results, run_id=run_id, tasks_file=tasks_file)
    report = {
        "summary": summary,
        "results": [asdict(result) for result in results],
    }
    results_dir.mkdir(parents=True, exist_ok=True)
    json_path = results_dir / f"memory_retrieval_{run_id}.json"
    md_path = results_dir / "memory_retrieval_latest.md"
    json_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    md_path.write_text(render_memory_retrieval_markdown(summary, results), encoding="utf-8")
    report["summary"]["json_report"] = str(json_path)
    report["summary"]["markdown_report"] = str(md_path)
    return report


def build_memory_items_from_payload(payload: dict[str, Any]) -> list[MemoryItem]:
    items: list[MemoryItem] = []
    timestamp = datetime.now().isoformat()
    for raw in payload.get("memory_items", []):
        if not isinstance(raw, dict):
            continue
        item_type = str(raw.get("type") or "fact")
        topic = str(raw.get("topic") or "")
        summary = str(raw.get("summary") or "")
        item_id = str(raw.get("id") or _memory_item_id(item_type, topic, summary))
        items.append(
            MemoryItem(
                id=item_id,
                timestamp=timestamp,
                updated_at=timestamp,
                source="benchmark",
                type=item_type,
                topic=topic,
                summary=summary,
                keywords=[str(value) for value in raw.get("keywords", [])],
                tags=[str(value) for value in raw.get("tags", [])],
                confidence=float(raw.get("confidence", 0.8) or 0.8),
                active=bool(raw.get("active", True)),
                embedding=list(raw.get("embedding") or []),
            )
        )
    return items


def first_relevant_rank(retrieved_ids: list[str], relevant_ids: list[str]) -> int | None:
    relevant = set(relevant_ids)
    for index, item_id in enumerate(retrieved_ids, start=1):
        if item_id in relevant:
            return index
    return None


def build_memory_retrieval_summary(
    results: list[MemoryRetrievalResult],
    *,
    run_id: str,
    tasks_file: Path,
) -> dict[str, Any]:
    total = len(results)
    hits = sum(1 for result in results if result.hit)
    reciprocal_sum = sum(result.reciprocal_rank for result in results)
    recall_by_k: dict[str, float] = {}
    for k in sorted({1, 3, 5, *[result.top_k for result in results]}):
        if k <= 0:
            continue
        hit_count = 0
        for result in results:
            if set(result.relevant_ids) & set(result.retrieved_ids[:k]):
                hit_count += 1
        recall_by_k[f"recall@{k}"] = round(hit_count / total, 4) if total else 0.0
    return {
        "run_id": run_id,
        "tasks_file": str(tasks_file),
        "total": total,
        "hits": hits,
        "misses": total - hits,
        "recall_at_task_k": round(hits / total, 4) if total else 0.0,
        "mrr": round(reciprocal_sum / total, 4) if total else 0.0,
        **recall_by_k,
    }


def render_memory_retrieval_markdown(
    summary: dict[str, Any],
    results: list[MemoryRetrievalResult],
) -> str:
    lines = [
        "# MiniAgent Memory Retrieval Benchmark",
        "",
        f"- Run ID: `{summary['run_id']}`",
        f"- Cases: {summary['total']}",
        f"- Hits: {summary['hits']}",
        f"- Misses: {summary['misses']}",
        f"- Recall@task_k: {summary['recall_at_task_k']:.2%}",
        f"- MRR: {summary['mrr']}",
        "",
        "## Recall",
        "",
    ]
    for key in sorted(k for k in summary if k.startswith("recall@")):
        lines.append(f"- `{key}`: {summary[key]:.2%}")
    lines.extend(["", "## Cases", ""])
    lines.append("| Case | Hit | Rank | Top K | Query |")
    lines.append("| --- | --- | ---: | ---: | --- |")
    for result in results:
        hit = "PASS" if result.hit else "MISS"
        rank = result.rank if result.rank is not None else "-"
        lines.append(f"| `{result.case_id}` | {hit} | {rank} | {result.top_k} | {result.query} |")
    return "\n".join(lines) + "\n"


def render_markdown_report(summary: dict[str, Any], results: list[BenchmarkResult]) -> str:
    lines = [
        "# MiniAgent Benchmark Report",
        "",
        f"- Run ID: `{summary['run_id']}`",
        f"- Tasks: {summary['total']}",
        f"- Passed: {summary['passed']}",
        f"- Failed: {summary['failed']}",
        f"- Success Rate: {summary['success_rate']:.2%}",
        f"- Total Tool Calls: {summary['total_tool_calls']}",
        f"- Avg Tool Calls: {summary['avg_tool_calls']}",
        f"- Avg Steps: {summary['avg_steps']}",
        "",
        "## Failure Types",
        "",
    ]
    if summary["failure_types"]:
        for name, count in sorted(summary["failure_types"].items()):
            lines.append(f"- `{name}`: {count}")
    else:
        lines.append("- None")
    lines.extend(["", "## Tasks", ""])
    lines.append("| Task | Category | Status | Steps | Tool Calls | Failure Types |")
    lines.append("| --- | --- | --- | ---: | ---: | --- |")
    for result in results:
        status = "PASS" if result.success else "FAIL"
        failures = ", ".join(f"`{item}`" for item in result.failure_types) or "-"
        lines.append(
            f"| `{result.task_id}` | {result.category} | {status} | "
            f"{result.steps} | {result.tool_call_count} | {failures} |"
        )
    return "\n".join(lines) + "\n"


def load_tasks(path: Path) -> list[BenchmarkTask]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    raw_tasks = payload.get("tasks", payload) if isinstance(payload, dict) else payload
    if not isinstance(raw_tasks, list):
        raise ValueError("Benchmark tasks file must contain a list or {'tasks': [...]} payload.")
    tasks: list[BenchmarkTask] = []
    for item in raw_tasks:
        if not isinstance(item, dict):
            continue
        tasks.append(BenchmarkTask(**item))
    return tasks


def ensure_default_benchmark_files(tasks_file: Path) -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    DEFAULT_RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    ensure_text_fixture(FIXTURES_DIR / "miniagent_notes.txt")
    ensure_xlsx_fixture(FIXTURES_DIR / "research_projects.xlsx")
    if not tasks_file.exists():
        tasks_file.parent.mkdir(parents=True, exist_ok=True)
        tasks_file.write_text(json.dumps(default_tasks_payload(), ensure_ascii=False, indent=2), encoding="utf-8")


def ensure_default_memory_retrieval_tasks(tasks_file: Path) -> None:
    tasks_file.parent.mkdir(parents=True, exist_ok=True)
    if tasks_file.exists():
        return
    tasks_file.write_text(
        json.dumps(default_memory_retrieval_payload(), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def ensure_text_fixture(path: Path) -> None:
    if path.exists():
        return
    path.write_text(
        "\n".join(
            [
                "MiniAgent 是一个 QQBot Agent 项目。",
                "核心能力包括文件接收、Office 文档处理、skill runtime 和长期记忆检索。",
                "当前重点是把能力系统化，而不是不断堆功能。",
            ]
        ),
        encoding="utf-8",
    )


def ensure_xlsx_fixture(path: Path) -> None:
    if path.exists():
        return
    try:
        from openpyxl import Workbook
    except ImportError:
        return
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Projects"
    sheet.append(["部门", "专业", "项目名称", "需求专业", "研究内容"])
    sheet.append(["实验室A", "软件工程", "AI 代码审查平台", "软件工程、人工智能", "构建大模型辅助代码审查系统"])
    sheet.append(["实验室B", "机械工程", "传感器结构设计", "机械工程", "硬件结构优化"])
    sheet.append(["实验室C", "软件工程", "数据治理工具", "软件工程", "企业数据清洗平台"])
    sheet.append(["实验室D", "人工智能", "车载 LLM 安全", "软件工程、网络安全", "研究 AI 大模型安全防御"])
    workbook.save(path)


def default_tasks_payload() -> dict[str, Any]:
    return {
        "tasks": [
            {
                "id": "basic_no_tool_math",
                "category": "chat",
                "prompt": "请只回答：2+2等于多少？",
                "expected_reply_contains": ["4"],
                "expected_max_tool_calls": 0,
            },
            {
                "id": "txt_attachment_summary",
                "category": "attachment",
                "prompt": "总结这个文本文件的核心内容，输出3条要点。",
                "attachments": [
                    {
                        "fixture": "miniagent_notes.txt",
                        "filename": "miniagent_notes.txt",
                        "content_type": "text/plain",
                    }
                ],
                "expected_reply_contains": ["MiniAgent"],
                "expected_tools_any": ["read_uploaded_file"],
            },
            {
                "id": "xlsx_filter_to_outbox",
                "category": "xlsx",
                "prompt": "筛选这个 Excel 中专业包含软件工程且研究内容包含 AI 或大模型的项目，保存为新的 Excel 文件。",
                "attachments": [
                    {
                        "fixture": "research_projects.xlsx",
                        "filename": "research_projects.xlsx",
                        "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    }
                ],
                "expected_tools_any": ["run_skill_script", "save_outbox_file"],
                "expected_outbox_suffixes": [".xlsx"],
                "expected_outbox_files": [
                    {
                        "suffix": ".xlsx",
                        "contains_all": ["AI 代码审查平台", "车载 LLM 安全"],
                        "not_contains": ["传感器结构设计", "数据治理工具"],
                        "min_rows": 3,
                        "max_rows": 4,
                    }
                ],
            },
            {
                "id": "generate_docx_report",
                "category": "docx",
                "prompt": "生成一份 Word 文档保存，标题为 MiniAgent Benchmark，内容包括一段项目简介。",
                "expected_tools_any": ["save_outbox_file", "run_skill_script"],
                "expected_outbox_suffixes": [".docx"],
                "expected_outbox_files": [
                    {
                        "suffix": ".docx",
                        "contains_all": ["MiniAgent Benchmark", "项目简介"],
                    }
                ],
            },
            {
                "id": "generate_pdf_report",
                "category": "pdf",
                "prompt": "生成一份 PDF 报告保存，标题为 MiniAgent Benchmark，内容包括一段项目简介。",
                "expected_tools_any": ["save_outbox_file", "run_skill_script"],
                "expected_outbox_suffixes": [".pdf"],
                "expected_outbox_files": [
                    {
                        "suffix": ".pdf",
                        "contains_all": ["MiniAgent Benchmark", "项目简介"],
                    }
                ],
            },
        ]
    }


def default_memory_retrieval_payload() -> dict[str, Any]:
    return {
        "memory_items": [
            {
                "id": "profile_identity",
                "type": "profile",
                "topic": "user identity",
                "summary": "用户姓名为小明，偏好简洁、专业、技术导向的中文交流。",
                "keywords": ["小明", "偏好", "技术中文"],
                "confidence": 0.95,
            },
            {
                "id": "project_root",
                "type": "project",
                "topic": "project environment",
                "summary": "MiniAgent 项目根目录位于 D:\\VScode\\project\\Agent\\AI_assistant，工作区目录为 workspace。",
                "keywords": ["项目根目录", "workspace", "AI_assistant"],
                "confidence": 0.95,
            },
            {
                "id": "xlsx_workflow",
                "type": "workflow",
                "topic": "xlsx file processing",
                "summary": "用户常用 Excel workflow：筛选研究内容包含 AI 且需求专业包含软件工程的课题，并输出新的 xlsx 文件。",
                "keywords": ["xlsx", "Excel", "AI", "软件工程", "筛选"],
                "confidence": 0.9,
            },
            {
                "id": "skill_runtime",
                "type": "tooling",
                "topic": "skill runtime",
                "summary": "当前 skill runtime 使用 run_skill_script 统一执行 workspace/skills/<skill>/scripts 下的 Python 脚本，并记录 skill_trace.jsonl。",
                "keywords": ["skill", "runtime", "run_skill_script", "skill_trace"],
                "confidence": 0.9,
            },
            {
                "id": "memory_retrieval",
                "type": "tooling",
                "topic": "memory retrieval",
                "summary": "长期记忆检索以 memory_store.jsonl 为唯一事实源，按 query embedding 召回 top_k，再结合关键词、topic 和 confidence 重排。",
                "keywords": ["memory_store", "embedding", "rerank", "top_k"],
                "confidence": 0.9,
            },
            {
                "id": "weather_skill",
                "type": "workflow",
                "topic": "weather skill",
                "summary": "天气查询应优先通过 weather/scripts/query_weather.py 执行，而不是直接拼 curl 或 PowerShell 请求。",
                "keywords": ["天气", "weather", "query_weather.py"],
                "confidence": 0.85,
            },
            {
                "id": "docx_generation",
                "type": "workflow",
                "topic": "docx generation",
                "summary": "简单 Word 文档可以通过 save_outbox_file 生成 docx，但这种方式不会产生 docx skill_script trace。",
                "keywords": ["docx", "Word", "save_outbox_file", "trace"],
                "confidence": 0.85,
            },
            {
                "id": "attachment_index_gap",
                "type": "project",
                "topic": "attachment visibility",
                "summary": "当前附件可见性依赖本轮附件和 session 中保留的附件记录，session 裁剪后旧附件可能需要独立 attachment index 才能稳定找回。",
                "keywords": ["附件", "session", "attachment index", "inbox"],
                "confidence": 0.85,
            },
        ],
        "queries": [
            {
                "id": "q_identity",
                "query": "你记得我的名字和沟通偏好吗？",
                "relevant_ids": ["profile_identity"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_project_root",
                "query": "当前项目根目录和 workspace 是什么关系？",
                "relevant_ids": ["project_root"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_xlsx_workflow",
                "query": "我之前处理 Excel 里 AI 和软件工程课题的流程是什么？",
                "relevant_ids": ["xlsx_workflow"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_skill_runtime",
                "query": "当前 skill runtime 是怎么执行脚本并记录日志的？",
                "relevant_ids": ["skill_runtime"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_memory_rerank",
                "query": "长期记忆是怎么做 embedding 检索和重排的？",
                "relevant_ids": ["memory_retrieval"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_weather",
                "query": "天气查询应该调用哪个 skill 脚本？",
                "relevant_ids": ["weather_skill"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_docx_trace",
                "query": "为什么生成 Word 后 trace 里可能没有 docx script success？",
                "relevant_ids": ["docx_generation"],
                "top_k": 4,
                "candidate_pool": 8,
            },
            {
                "id": "q_attachment_visibility",
                "query": "为什么 session 裁剪后可能找不到以前上传的附件？",
                "relevant_ids": ["attachment_index_gap"],
                "top_k": 4,
                "candidate_pool": 8,
            },
        ],
    }


def dedupe(values: list[str]) -> list[str]:
    seen: set[str] = set()
    results: list[str] = []
    for value in values:
        if value in seen:
            continue
        seen.add(value)
        results.append(value)
    return results


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Run MiniAgent benchmark task set.")
    parser.add_argument(
        "suite",
        nargs="?",
        default="agent",
        choices=["agent", "memory"],
        help="Benchmark suite to run. Default: agent.",
    )
    parser.add_argument("--tasks", type=Path, default=DEFAULT_TASKS_FILE, help="Benchmark task JSON file.")
    parser.add_argument("--limit", type=int, default=None, help="Run only the first N tasks.")
    parser.add_argument("--delay", type=float, default=3.0, help="Seconds to wait between agent eval tasks. Use 0 to disable.")
    parser.add_argument("--json", action="store_true", help="Print full JSON report.")
    return parser


async def async_main(argv: list[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    tasks_file = args.tasks
    if args.suite == "memory" and tasks_file == DEFAULT_TASKS_FILE:
        tasks_file = DEFAULT_MEMORY_TASKS_FILE
    if args.suite == "memory":
        report = await run_memory_retrieval_benchmark(tasks_file=tasks_file)
    else:
        report = await run_benchmark(
            tasks_file=tasks_file,
            limit=args.limit,
            delay_seconds=max(0.0, args.delay),
        )
    if args.json:
        print(json.dumps(report, ensure_ascii=False, indent=2))
    else:
        summary = report["summary"]
        if args.suite == "memory":
            print("\nMemory Retrieval Benchmark Summary")
            print(f"Run ID: {summary['run_id']}")
            print(f"Cases: {summary['total']}")
            print(f"Hits: {summary['hits']}")
            print(f"Misses: {summary['misses']}")
            print(f"Recall@task_k: {summary['recall_at_task_k']:.2%}")
            print(f"MRR: {summary['mrr']}")
            for key in sorted(k for k in summary if k.startswith("recall@")):
                print(f"{key}: {summary[key]:.2%}")
            print(f"JSON Report: {summary['json_report']}")
            print(f"Markdown Report: {summary['markdown_report']}")
            return 1 if summary["misses"] else 0

        print("\nBenchmark Summary")
        print(f"Run ID: {summary['run_id']}")
        print(f"Tasks: {summary['total']}")
        print(f"Passed: {summary['passed']}")
        print(f"Failed: {summary['failed']}")
        print(f"Success Rate: {summary['success_rate']:.2%}")
        print(f"Total Tool Calls: {summary['total_tool_calls']}")
        print(f"Avg Tool Calls: {summary['avg_tool_calls']}")
        print(f"Avg Steps: {summary['avg_steps']}")
        print(f"Failure Types: {summary['failure_types'] or {}}")
        print(f"JSON Report: {summary['json_report']}")
        print(f"Markdown Report: {summary['markdown_report']}")
    return 1 if report["summary"]["failed"] else 0


def main(argv: list[str] | None = None) -> int:
    return asyncio.run(async_main(argv))
